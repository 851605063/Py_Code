import logging
import os
import re
import datetime
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as WD_ALIGN_PARAGRAPH

logger = logging.getLogger()
logger.setLevel(logging.INFO)
basic_url = "http://paper.people.com.cn/rmrb/html/"


# dict={'h3':'xxx','h1':'xxx','h2':'xxx','secP':secP,'articleDate':articleDate,'content':[]}
def saveNews(dict, date):
    path = 'rmrb111/' + date
    if os.path.exists(path) is False:
        os.makedirs(path)
    doc = Document()
    doc.sections[0].right_margin=Cm(6) #右页边距
    normal_style = doc.styles['Normal']
    normal_style.font.name = '宋体'
    normal_style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    # 首行缩进
    para_format = normal_style.paragraph_format
    para_format.first_line_indent = Cm(0.74)
    if dict['h3'] != 'None':
        h3 = doc.add_heading(dict['h3'], level=3)  # 插入h3
        h3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1 = doc.add_heading(dict['h1'], level=0)  # 插入h1
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if dict['h2'] != 'None':
        h2 = doc.add_heading(dict['h2'], level=1)  # 插入h2
        h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # secP=doc.add_paragraph(dict['secP']) #插入作者
    # secP.alignment = WD_ALIGN_PARAGRAPH.CENTER
    articleDate = doc.add_paragraph(dict['articleDate']) #插入日期
    articleDate.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #插入段落
    for p in dict['contents']:
        doc.add_paragraph(p)
    docname=path+'/'+dict['h1']+'.docx'
    doc.save(docname)
def getCommentHref(date):
    url = basic_url + "{}/nbs.D110000renmrb_01.htm".format(date)
    r = requests.get(url)
    soup = BeautifulSoup(r.text, features="lxml")
    list = soup.find_all("a", href=re.compile('nbs.D110000renmrb.*'))
    comment_a = ''
    for a in list:
        if str(a.string).find('评论') != -1:
            comment_a = str(a.string)
    if comment_a == '':
        return None
    s = comment_a.split('版：')
    url = url.replace('_01.', '_' + s[0] + '.')
    return url


def getCommentNewsurl(url, date):
    r = requests.get(url)
    soup = BeautifulSoup(r.text, features="lxml")
    ul = soup.find_all('ul', class_="news-list")
    list = re.findall('<a href="(.*)"', str(ul))
    content_urls = []
    for l in list:
        content_url = basic_url + date + '/' + l
        content_urls.append(content_url)
    return content_urls
def getNewsContent_OneDay(date):
    # 获取评论版面链接
    url = getCommentHref(date)
    if url == None:
        return
    logging.info('[' + date + ']' + url)
    # 获取评论版面新闻链接
    list = getCommentNewsurl(url, date)
    # 遍历链接获取内容存储
    for l in list:
        dict = getNewsbyList(l, date)
        saveNews(dict, date)
def getNewsContent_More(days):
    days = int(days) - 1
    while days >= 0:
        date = (datetime.date.today() - datetime.timedelta(days)).__format__('%Y-%m/%d')
        days = days - 1
        getNewsContent_OneDay(date)
def getNewsbyList(url, date):
    r = requests.get(url)
    soup = BeautifulSoup(r.text, 'lxml')
    articleDiv = soup.find('div', class_='article')
    h3 = str(articleDiv.find('h3').string)
    h1 = str(articleDiv.find('h1').string)
    h2 = str(articleDiv.find('h2').string)
    secP=str(articleDiv.find('p',class_='sec').string)
    articleDate=re.sub('\n|\r| +','',str(articleDiv.find('span', class_='date').string))
    print(articleDate)
    article = {'h3': h3, 'h1': h1, 'h2': h2,'secP':secP,'articleDate':articleDate}
    articleContentDiv = soup.find('div', id='articleContent')
    articleContentPgs = articleContentDiv.find_all('p')
    listP = []
    for p in articleContentPgs:
        item=re.sub(u'\u3000+',u'',str(p.string))
        if item=='':
            continue
        listP.append(item)
        #str(p.string).replace(u'\u3000', u'')
    article['contents'] = listP
    return article
if __name__ == "__main__":
    print('请输入需要获取的天数：')
    days = input()
    getNewsContent_More(days)
